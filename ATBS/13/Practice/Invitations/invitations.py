#! python3

import docx

# Write a program that would generate a Word document with 
# custom invitations

###### doesn work. needs to be fixed #####

#open the doc
doc = docx.Document('invitation.docx')
file = open('guests.txt', 'r')

for line in file.readlines():
    name = line.strip('\n')
    doc.add_paragraph('Dear')
    doc.add_paragraph(name)
    doc.add_paragraph('Please come to the party')
    doc.add_paragraph('Feb 12')
    doc.add_paragraph('19:00')

    # run = doc.paragraphs[len(doc.paragraphs)].add_run()
    # run.add_break(WD_BREAK.PAGE)
    
doc.save('invitations_bulk.docx')