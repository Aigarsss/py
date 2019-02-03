#! python3

# installed pyPDF2 and Python-Docx

'''
# impoting and opening
import PyPDF2

pdfFileObj = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
print(pdfReader.numPages)
pageObj = pdfReader.getPage(0)
print(pageObj.extractText())

# encryption

import PyPDF2
pdfReader = PyPDF2.PdfFileReader(open('encrypted.pdf', 'rb'))
print(pdfReader.isEncrypted) # true
#print(pdfReader.getPage(0)) # gives and error
pdfReader.decrypt('rosebud')
pageObj = pdfReader.getPage(0)
print(pageObj) # returns some weird values

# creating pdfs
import PyPDF2
pdf1File = open('meetingminutes.pdf', 'rb')
pdf2File = open('meetingminutes2.pdf', 'rb')
pdf1Reader = PyPDF2.PdfFileReader(pdf1File)
pdf2Reader = PyPDF2.PdfFileReader(pdf2File)
pdfWriter = PyPDF2.PdfFileWriter()

for pageNum in range(pdf1Reader.numPages):
    pageObj = pdf1Reader.getPage(pageNum)
    pdfWriter.addPage(pageObj)

for pageNum in range(pdf2Reader.numPages):
    pageObj = pdf2Reader.getPage(pageNum)
    pdfWriter.addPage(pageObj)

pdfOutputFile = open('combinedminutes.pdf', 'wb')
pdfWriter.write(pdfOutputFile)
pdfOutputFile.close()
pdf1File.close()
pdf2File.close()


# rotating pages
import PyPDF2
minutesFile = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(minutesFile)
page = pdfReader.getPage(0)
page.rotateClockwise(90)

pdfWriter = PyPDF2.PdfFileWriter()
pdfWriter.addPage(page)
resultPdfFile = open('RotatedPage.pdf', 'wb')
pdfWriter.write(resultPdfFile)
resultPdfFile.close()
minutesFile.close()

# overlay and watermarking
import PyPDF2
minutesFile = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(minutesFile)
minutesFirstPage = pdfReader.getPage(0)
pdfWaterMarkReader = PyPDF2.PdfFileReader(open('watermark.pdf', 'rb'))
minutesFirstPage.mergePage(pdfWaterMarkReader.getPage(0))

pdfWriter = PyPDF2.PdfFileWriter()
pdfWriter.addPage(minutesFirstPage)

for pageNum in range(1, pdfReader.numPages):
    pageObj = pdfReader.getPage(pageNum)
    pdfWriter.addPage(pageObj)

resultPdfFile = open('watermarkedCover.pdf', 'wb')
pdfWriter.write(resultPdfFile)
minutesFile.close()
resultPdfFile.close()


# encryption
import PyPDF2
pdfFile = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFile)
pdfWriter = PyPDF2.PdfFileWriter()
for pageNum in range(pdfReader.numPages):
    pdfWriter.addPage(pdfReader.getPage(pageNum))

pdfWriter.encrypt('swordfish')
resultPdf = open('encryptedminutes.pdf', 'wb')
pdfWriter.write(resultPdf)
pdfFile.close()
resultPdf.close()


#### Word Documents ####
import docx
doc = docx.Document('demo.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[1].text) # A plain paragraph with some bold and some italic
print(len(doc.paragraphs[1].runs)) # 5 runs
print(doc.paragraphs[1].runs[0].text) # A plain paragraph with

# getting the full info, without fortmats
import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('demo.docx'))

# changing styles
import docx

doc = docx.Document('demo.docx')
print(doc.paragraphs[0].text) # Document Title
print(doc.paragraphs[0].style) # _ParagraphStyle('Title') id: 52507792

doc.paragraphs[0].style = 'Normal'
print(doc.paragraphs[1].text)

print(doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)

doc.paragraphs[1].runs[0] = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('restyled.docx')

# writing docs one paragraph
import docx
doc = docx.Document()
doc.add_paragraph('Hello world')
doc.save('helloworld.docx')

# writing docs multiple paragraphs
import docx
doc = docx.Document()
doc.add_paragraph('Hello world')

paraObj1 = doc.add_paragraph('This is the second')
paraObj2 = doc.add_paragraph('This is the third')
paraObj1.add_run(' this text is being dded to the second paragraph')
doc.save('multipara.docx')

# adding headings
import docx
doc = docx.Document()
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)
doc.save('headings.docx')

# adding images
import docx
doc = docx.Document()
doc.add_picture('image.png', width = docx.shared.Cm(4), height = docx.shared.Cm(4)) # params optional

'''

# line and page breaks
# import docx
# doc = docx.Document()
# doc.add_paragraph('This is on the first page')
# doc.paragraphs[0].runs[0].add_break(docx.text.WD_BREAK.PAGE) # doesnt work...
# doc.add_paragraph('This is in the second')
# doc.save('two_page.docx')

